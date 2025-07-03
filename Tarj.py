import fitz
import re
from selenium.webdriver.common.by import By
from docx import Document
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from tkinter import Tk, Button, Label
from datetime import datetime
from tkinter import Toplevel, Text, Scrollbar, RIGHT, Y, END
import pandas as pd
import os
from tkinter import Entry, StringVar, filedialog
from tkinter import Tk, Button, Label, ttk
from tkinter import simpledialog
from tkinter import messagebox
from tkinter import PhotoImage
from tkinter import Tk, Label
from PIL import Image, ImageTk 

import json
config_path = "config_ikarus.json"
 
# ───────────────────────────────
# HISTÓRICO
# ───────────────────────────────
def log_taj(acao):
    with open("historico_taj.log", "a", encoding="utf-8") as log:
        log.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {acao}\n")

def iniciar_interface():
    janela = Tk()
    janela.title("TAJ - Assistente Digital")
    janela.geometry("400x350")
    janela.configure(bg="#c0da2f")  

    imagem = Image.open("logo_PrefSP_sem fundo_vertical__branco_monocromático.png")
    imagem = imagem.resize((120, 120))  
    logo = ImageTk.PhotoImage(imagem)
    Label(janela, image=logo, bg="#0A3979").pack(pady=10)
    janela.logo = logo  


    Label(janela, text="TAJ-SP - Assistente Digital", font=("Arial", 14), bg="#0A3979", fg="white").pack(pady=5)

 
    estilo = ttk.Style()
    estilo.configure("Custom.TButton",
                     font=("Segoe UI", 11, "bold"),
                     foreground="white",
                     background="#0074c1",
                     padding=6)
    estilo.map("Custom.TButton",
               background=[("active", "#f7931e")])
 
    janela.title("🦅 TAJ -SP - Assistente Digital")
    janela.geometry("400x350")
    janela.configure(bg="#0A3979")

   
    Button(janela, text="📂 Ver Histórico de Ações", width=30, command=ver_historico).pack(pady=8)
    Button(janela, text= "🔐 Tarjar Dados Sensíveis em PDF", width=30, command=tarjar_pdf).pack(pady=8)
    Button(janela, text="📝 Tarjar Dados Sensíveis em Word", width=30, command=tarjar_docx).pack(pady=8)
    Button(janela, text="📊 Tarjar Dados Sensíveis em CSV", width=30, command=tarjar_csv).pack(pady=8)
    Button(janela, text="❌ Sair", width=30, command=janela.destroy).pack(pady=12)

 
    janela.mainloop()
 
def ver_historico():
    historico_janela = Toplevel()
    historico_janela.title("📂 Histórico de Ações do Àguia")
    historico_janela.geometry("600x400")
    scrollbar = Scrollbar(historico_janela)
    scrollbar.pack(side=RIGHT, fill=Y)
    texto = Text(historico_janela, wrap="word", yscrollcommand=scrollbar.set)
    texto.pack(expand=True, fill="both")
 
    try:
        with open("historico_ikarus.log", "r", encoding="utf-8") as log_file:
            texto.insert(END, log_file.read())
    except FileNotFoundError:
        texto.insert(END, "⚠️ Nenhum histórico encontrado.")
 
    scrollbar.config(command=texto.yview)
  
      
def tarjar_docx():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione o arquivo Word", filetypes=[("Word Documents", "*.docx")]
    )
    if not caminho_arquivo:
        return

    padroes = {
        "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
        "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
        "Telefone": r"\b\(?\d{2}\)?\s?\d{4,5}-\d{4}\b",
        "E-mail": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
        "Senha": r"\bsenha\s*[:=]?\s*\S+",
        "Processo CNJ": r"\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b",
        "CEP": r"\b\d{5}-\d{3}\b",
        "Cartão de Crédito": r"\b(?:\d[ -]*?){13,16}\b",
        "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
        "Passaporte": r"\b[A-Z]{1}\d{7}\b",
    }

    doc = Document(caminho_arquivo)
    total_ocultados = 0

    def substituir(texto):
        nonlocal total_ocultados
        for padrao in padroes.values():
            if re.search(padrao, texto):
                texto = re.sub(padrao, "[TARJADO]", texto)
                total_ocultados += 1
        return texto

    for paragrafo in doc.paragraphs:
        paragrafo.text = substituir(paragrafo.text)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                celula.text = substituir(celula.text)

    if total_ocultados:
        novo_nome = caminho_arquivo.replace(".docx", "_TARJADO.docx")
        doc.save(novo_nome)
        messagebox.showinfo(
            "Sucesso",
            f"{total_ocultados} dados sensíveis foram tarjados.\nArquivo salvo como:\n{novo_nome}",
        )
        log_taj("Dados sensíveis tarjados em Word.")
    else:
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível encontrado em Word.")
            
def tarjar_pdf():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione o PDF", filetypes=[("PDF Files", "*.pdf")]
    )
    if not caminho_arquivo:
        return

    padroes = {
        "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
        "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
        "Telefone": r"\b\(?\d{2}\)?\s?\d{4,5}-\d{4}\b",
        "E-mail": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
        "Senha": r"\bsenha\s*[:=]?\s*\S+",
        "Processo CNJ": r"\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b",
        "CEP": r"\b\d{5}-\d{3}\b",
        "Cartão de Crédito": r"\b(?:\d[ -]*?){13,16}\b",
        "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
        "Passaporte": r"\b[A-Z]{1}\d{7}\b",
    }

    doc = fitz.open(caminho_arquivo)
    total_ocultados = 0
    has_redaction = hasattr(doc, "apply_redactions")

    for page in doc:
        texto_pagina = page.get_text()
        for padrao in padroes.values():
            for ocorrencia in re.finditer(padrao, texto_pagina, re.IGNORECASE):
                texto_encontrado = ocorrencia.group()
                areas = page.search_for(texto_encontrado)
                for area in areas:
                    if has_redaction:
                        page.add_redact_annot(
                            area, fill=(0, 0, 0), text="[TARJADO]", align=1
                        )
                    else:
                        # fallback para versões antigas
                        page.draw_rect(area, color=(0, 0, 0), fill=(0, 0, 0))
                    total_ocultados += 1

    if total_ocultados:
        if has_redaction:
            doc.apply_redactions()
            msg = "Redação completa aplicada. Dados removidos."
        else:
            msg = ("ATENÇÃO: seu PyMuPDF não suporta redactions.\n"
                   "O texto ainda pode ser selecionado por baixo.")
        novo_nome = caminho_arquivo.replace(".pdf", "_TARJADO.pdf")
        doc.save(novo_nome)
        doc.close()
        messagebox.showinfo(
            "Sucesso",
            f"{total_ocultados} dados sensíveis foram tarjados.\n{msg}\nArquivo salvo como:\n{novo_nome}",
        )
        log_taj("Dados sensíveis tarjados em PDF.")
    else:
        doc.close()
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível encontrado em PDF.")



def substituir(texto, padroes, total_ocultados):
    for padrao in padroes.values():
        if re.search(padrao, texto):
            texto = re.sub(padrao, "[TARJADO]", texto)
            total_ocultados[0] += 1
    return texto

def pre_visualizar_docx():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione o arquivo Word", filetypes=[("Word Documents", "*.docx")]
    )
    if not caminho_arquivo:
        return

    padroes = {
        "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
        "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
        "Telefone": r"\b\(?\d{2}\)?\s?\d{4,5}-\d{4}\b",
        "E-mail": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
        "Senha": r"\bsenha\s*[:=]?\s*\S+",
        "Processo CNJ": r"\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b",
        "CEP": r"\b\d{5}-\d{3}\b",
        "Cartão de Crédito": r"\b(?:\d[ -]*?){13,16}\b",
        "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
        "Passaporte": r"\b[A-Z]{1}\d{7}\b",
    }

    doc = Document(caminho_arquivo)
    total_ocultados = [0]
    conteudo_preview = ""

    for paragrafo in doc.paragraphs:
        paragrafo.text = substituir(paragrafo.text, padroes, total_ocultados)
        conteudo_preview += paragrafo.text + "\n"

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                celula.text = substituir(celula.text, padroes, total_ocultados)
                conteudo_preview += celula.text + "\n"

    if total_ocultados[0] == 0:
        messagebox.showinfo("Nada encontrado", "Nenhum dado sensível detectado.")
        return

    preview = Toplevel()
    preview.title("Pré-visualização do documento tarjado")
    preview.geometry("600x500")

    text_widget = Text(preview, wrap="word")
    text_widget.insert(END, conteudo_preview)
    text_widget.pack(expand=True, fill="both")

    def salvar():
        novo_nome = caminho_arquivo.replace(".docx", "_TARJADO.docx")
        doc.save(novo_nome)
        messagebox.showinfo("Sucesso", f"Documento salvo como:\n{novo_nome}")
        preview.destroy()

    Button(preview, text="Salvar documento", command=salvar).pack(pady=10)




def tarjar_csv():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione o arquivo CSV", filetypes=[("CSV Files", "*.csv")]
    )
    if not caminho_arquivo:
        return

    padroes = {
        "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
        "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
        "Telefone": r"\b\(?\d{2}\)?\s?\d{4,5}-\d{4}\b",
        "E-mail": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
        "Senha": r"\bsenha\s*[:=]?\s*\S+",
        "Processo CNJ": r"\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b",
        "CEP": r"\b\d{5}-\d{3}\b",
        "Cartão de Crédito": r"\b(?:\d[ -]*?){13,16}\b",
        "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
        "Passaporte": r"\b[A-Z]{1}\d{7}\b",
    }

    try:
        df = pd.read_csv(caminho_arquivo, dtype=str)  # Evita erros com tipos mistos
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao ler o CSV:\n{e}")
        return

    total_ocultados = 0

    for coluna in df.columns:
        df[coluna] = df[coluna].astype(str).apply(
            lambda x: ocultar_dados(x, padroes, total_ocultados_ref := [0])
        )
        total_ocultados += total_ocultados_ref[0]

    if total_ocultados:
        novo_nome = caminho_arquivo.replace(".csv", "_TARJADO.csv")
        df.to_csv(novo_nome, index=False)
        messagebox.showinfo(
            "Sucesso",
            f"{total_ocultados} dados sensíveis foram tarjados.\nArquivo salvo como:\n{novo_nome}",
        )
        log_taj("Dados sensíveis tarjados em CSV.")
    else:
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível encontrado em CSV.")

def ocultar_dados(texto, padroes, total_ocultados):
    for padrao in padroes.values():
        if re.search(padrao, texto):
            texto = re.sub(padrao, "[TARJADO]", texto)
            total_ocultados[0] += 1
    return texto
  
if __name__ == "__main__":
    iniciar_interface()