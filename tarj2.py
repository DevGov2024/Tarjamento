import os
import re
import glob
import csv
import fitz
import pandas as pd
from datetime import datetime
from tkinter import Tk, Toplevel, Text, Scrollbar, Label, Button, filedialog, messagebox, RIGHT, Y, END
from tkinter import ttk
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document

# ------------------- PADRÕES REGEX --------------------
padroes = {
    "CPF": r"\b\d{3}\.\d{3}\.\d{3}-\d{2}\b",
    "CNPJ": r"\b\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}\b",
    "Telefone": r"\b\(?\d{2}\)?\s?\d{4,5}-\d{4}\b",
    "E-mail": r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b",
    "Senha": r"\bsenha\s*[:=]?\s*\S+",
    "Processo CNJ": r"\b\d{7}-\d{2}\.\d{4}\.\d{1}\.\d{2}\.\d{4}\b",
    "CEP": r"\b\d{5}-\d{3}\b",
    "Cartão de Crédito": r"\b(?:\d[ -]*?){13,16}\b",
    "RG": r"\b\d{2}\.\d{3}\.\d{3}-\d{1}\b",
    "Passaporte": r"\b[A-Z]{1}\d{7}\b",
}

# ------------------- FUNÇÕES DE APOIO --------------------
def log_taj(acao):
    with open("historico_taj.log", "a", encoding="utf-8") as log:
        log.write(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {acao}\n")

def substituir(texto, padroes, total_ocultados):
    for padrao in padroes.values():
        if re.search(padrao, texto):
            texto = re.sub(padrao, "[TARJADO]", texto)
            total_ocultados[0] += 1
    return texto

def ocultar_dados(texto, padroes, total_ocultados):
    for padrao in padroes.values():
        if re.search(padrao, texto):
            texto = re.sub(padrao, "[TARJADO]", texto)
            total_ocultados[0] += 1
    return texto

relatorio_sessao = []

def adicionar_ao_relatorio(arquivo, tipo, quantidade):
    relatorio_sessao.append([datetime.now().strftime("%Y-%m-%d %H:%M:%S"), arquivo, tipo, quantidade])

def exportar_relatorio():
    if not relatorio_sessao:
        messagebox.showinfo("Relatório", "Nenhuma ação foi registrada nesta sessão.")
        return

    caminho = filedialog.asksaveasfilename(
        defaultextension=".csv",
        filetypes=[("CSV", "*.csv"), ("Texto", "*.txt")],
        title="Salvar relatório da sessão"
    )
    if not caminho:
        return

    try:
        with open(caminho, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(["Data/Hora", "Arquivo", "Tipo", "Qtd Dados Tarjados"])
            writer.writerows(relatorio_sessao)
        messagebox.showinfo("Sucesso", f"Relatório salvo em:\n{caminho}")
    except Exception as e:
        messagebox.showerror("Erro", f"Falha ao salvar relatório:\n{e}")

# ------------------- TARJAMENTO PDF --------------------
def tarjar_pdf():
    caminho_arquivo = filedialog.askopenfilename(title="Selecione o PDF", filetypes=[("PDF Files", "*.pdf")])
    if not caminho_arquivo:
        return

    try:
        doc = fitz.open(caminho_arquivo)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir o PDF:\n{e}")
        return

    if doc.page_count == 0:
        messagebox.showwarning("PDF Vazio", "O arquivo PDF não possui páginas.")
        doc.close()
        return

    if doc.needs_pass:
        messagebox.showerror("Protegido por Senha", "Este PDF está protegido por senha e não pode ser processado.")
        doc.close()
        return

    total_ocultados = 0
    has_redaction = hasattr(doc, "apply_redactions")

    padroes_escolhidos = selecionar_padroes()
    if not padroes_escolhidos:
       messagebox.showinfo("Cancelado", "Nenhum padrão foi selecionado.")
       return

    for page in doc:
        texto_pagina = page.get_text()
        for padrao in padroes_escolhidos.values():
            for ocorrencia in re.finditer(padrao, texto_pagina, re.IGNORECASE):
                texto_encontrado = ocorrencia.group()
                areas = page.search_for(texto_encontrado)
                for area in areas:
                    if has_redaction:
                       page.add_redact_annot(area, fill=(0, 0, 0), text="[TARJADO]", align=1)
                    else:
                       page.draw_rect(area, color=(0, 0, 0), fill=(0, 0, 0))
                    total_ocultados += 1

    if total_ocultados:
        if has_redaction:
            doc.apply_redactions()
            msg = "Redação aplicada."
        else:
            msg = "Texto ainda pode ser selecionado por baixo."
        novo_nome = caminho_arquivo.replace(".pdf", "_TARJADO.pdf")
        doc.save(novo_nome)
        messagebox.showinfo("Sucesso", f"{total_ocultados} dados tarjados.\n{msg}\nSalvo em:\n{novo_nome}")
        log_taj("PDF tarjado.")
        adicionar_ao_relatorio(caminho_arquivo, "PDF", total_ocultados)
    else:
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível em PDF.")

    doc.close()

# ------------------- TARJAMENTO WORD --------------------
def tarjar_docx():
    caminho_arquivo = filedialog.askopenfilename(title="Selecione o Word", filetypes=[("Word", "*.docx")])
    if not caminho_arquivo:
        return

    doc = Document(caminho_arquivo)
    total_ocultados = [0]

    for p in doc.paragraphs:
        p.text = substituir(p.text, padroes, total_ocultados)
    for t in doc.tables:
        for linha in t.rows:
            for cel in linha.cells:
                cel.text = substituir(cel.text, padroes, total_ocultados)

    if total_ocultados[0]:
        novo_nome = caminho_arquivo.replace(".docx", "_TARJADO.docx")
        doc.save(novo_nome)
        messagebox.showinfo("Sucesso", f"{total_ocultados[0]} dados tarjados.\nSalvo como:\n{novo_nome}")
        log_taj("Word tarjado.")
        adicionar_ao_relatorio(caminho_arquivo, "Word", total_ocultados[0])
    else:
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível em Word.")

# ------------------- TARJAMENTO CSV --------------------
def tarjar_csv():
    caminho_arquivo = filedialog.askopenfilename(title="Selecione o CSV", filetypes=[("CSV", "*.csv")])
    if not caminho_arquivo:
        return

    try:
        df = pd.read_csv(caminho_arquivo, dtype=str)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao ler CSV:\n{e}")
        return

    total_ocultados = 0

    for coluna in df.columns:
        df[coluna] = df[coluna].astype(str).apply(
            lambda x: ocultar_dados(x, padroes, ref := [0])
        )
        total_ocultados += ref[0]

    if total_ocultados:
        novo_nome = caminho_arquivo.replace(".csv", "_TARJADO.csv")
        df.to_csv(novo_nome, index=False)
        messagebox.showinfo("Sucesso", f"{total_ocultados} dados tarjados.\nSalvo como:\n{novo_nome}")
        log_taj("CSV tarjado.")
        adicionar_ao_relatorio(caminho_arquivo, "CSV", total_ocultados)
    else:
        messagebox.showinfo("Nada Encontrado", "Nenhum dado sensível encontrado.")
        log_taj("Nenhum dado sensível em CSV.")

# ------------------- MODO BATCH --------------------


# ------------------- INTERFACE --------------------
def iniciar_interface():
    janela = Tk()
    janela.title("TARJ-SP - Assistente Digital")
    janela.geometry("400x500")
    janela.configure(bg="#0A3979")

    imagem = Image.open("logo_PrefSP_sem fundo_vertical__branco_monocromático.png")
    imagem = imagem.resize((120, 120))  
    logo = ImageTk.PhotoImage(imagem)
    Label(janela, image=logo, bg="#0A3979").pack(pady=10)
    janela.logo = logo  
    Label(janela, text="TARJ-SP - Assistente Digital", font=("Arial", 14), bg="#0A3979", fg="white").pack(pady=5)

    def criar_botao(texto, comando):
        return Button(
            janela,
            text=texto,
            width=35,
            bg="#f7931e",
            fg="white",
            activebackground="#ffa733",
            font=("Segoe UI", 10, "bold"),
            command=comando
        )

    
    criar_botao("📂 Ver Histórico de Ações", ver_historico).pack(pady=5)
    criar_botao("🔐 Tarjar PDF", tarjar_pdf).pack(pady=5)
    criar_botao("📝 Tarjar Word", tarjar_docx).pack(pady=5)
    criar_botao("📊 Tarjar CSV", tarjar_csv).pack(pady=5)
    criar_botao("👁️ Pré-visualizar PDF", visualizar_pdf).pack(pady=3)
    criar_botao("👁️ Pré-visualizar Word", visualizar_docx).pack(pady=3)
    criar_botao("👁️ Pré-visualizar CSV", visualizar_csv).pack(pady=3)
    criar_botao("❌ Sair", janela.destroy).pack(pady=10)

    
    janela.mainloop()

def ver_historico():
    historico_janela = Toplevel()
    historico_janela.title("📂 Histórico de Ações")
    historico_janela.geometry("600x400")
    scrollbar = Scrollbar(historico_janela)
    scrollbar.pack(side=RIGHT, fill=Y)
    texto = Text(historico_janela, wrap="word", yscrollcommand=scrollbar.set)
    texto.pack(expand=True, fill="both")

    try:
        with open("historico_taj.log", "r", encoding="utf-8") as log_file:
            texto.insert(END, log_file.read())
    except FileNotFoundError:
        texto.insert(END, "⚠️ Nenhum histórico encontrado.")

    scrollbar.config(command=texto.yview)

def visualizar_pdf():
    caminho = filedialog.askopenfilename(title="Selecione um PDF", filetypes=[("PDF", "*.pdf")])
    if not caminho:
        return
    try:
        doc = fitz.open(caminho)
        if doc.page_count == 0:
            messagebox.showwarning("PDF Vazio", "O arquivo não possui páginas.")
            return
        page = doc[0]
        pix = page.get_pixmap()
        image_path = "preview_temp.png"
        pix.save(image_path)
        doc.close()
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir PDF:\n{e}")
        return

    janela = Toplevel()
    janela.title("Pré-visualização PDF")
    imagem = Image.open(image_path)
    imagem_tk = ImageTk.PhotoImage(imagem)
    label = Label(janela, image=imagem_tk)
    label.image = imagem_tk
    label.pack()

def visualizar_docx():
    caminho = filedialog.askopenfilename(title="Selecione um Word", filetypes=[("Word", "*.docx")])
    if not caminho:
        return
    try:
        doc = Document(caminho)
        parags = [p.text for p in doc.paragraphs if p.text.strip()]
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir DOCX:\n{e}")
        return

    janela = Toplevel()
    janela.title("Pré-visualização Word")
    texto = Text(janela, wrap="word")
    texto.pack(expand=True, fill="both")
    texto.insert(END, "\n".join(parags[:20]))  # mostra até 20 parágrafos

def visualizar_csv():
    caminho = filedialog.askopenfilename(title="Selecione um CSV", filetypes=[("CSV", "*.csv")])
    if not caminho:
        return
    try:
        df = pd.read_csv(caminho)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir CSV:\n{e}")
        return

    janela = Toplevel()
    janela.title("Pré-visualização CSV")
    texto = Text(janela, wrap="none")
    texto.pack(expand=True, fill="both")
    texto.insert(END, df.head(20).to_string(index=False))


def selecionar_padroes():
    padroes_escolhidos = {}

    def confirmar():
        for chave, var in check_vars.items():
            if var.get():
                padroes_escolhidos[chave] = padroes[chave]
        janela.destroy()

    janela = Toplevel()
    janela.title("Selecionar Dados a Tarjar")
    Label(janela, text="Escolha quais dados devem ser tarjados:", font=("Arial", 10, "bold")).pack(pady=10)

    check_vars = {}
    for chave in padroes:
        var = tk.IntVar(value=1)  # Todos marcados por padrão
        chk = tk.Checkbutton(janela, text=chave, variable=var)
        chk.pack(anchor="w")
        check_vars[chave] = var

    Button(janela, text="Confirmar", command=confirmar).pack(pady=10)
    janela.wait_window()

    return padroes_escolhidos




# ------------------- EXECUÇÃO --------------------
if __name__ == "__main__":
    iniciar_interface()
